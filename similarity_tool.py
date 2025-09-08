"""
Document Similarity Analysis Tool
Detects and clusters similar paragraphs in documents using sentence embeddings
Windows & macOS compatible (with Linux fallback)
"""

from openai import OpenAI
import os
import sys
import docx
import fitz  # PyMuPDF for PDFs
from sklearn.metrics.pairwise import cosine_distances
from sklearn.cluster import DBSCAN
from sentence_transformers import SentenceTransformer
import numpy as np
import csv
import logging
from typing import List, Tuple, Dict, Optional
from dataclasses import dataclass
import json
from pathlib import Path
import platform
import hashlib

# --------------------
# CONFIGURATION
# --------------------
@dataclass
class Config:
    """Configuration settings for similarity analysis"""
    api_key: Optional[str] = os.getenv("OPENAI_API_KEY")
    # Jeśli analizujesz sporo polskich tekstów, rozważ:
    # "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
    embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2"
    chunk_size: int = 20
    eps: float = 0.3               # DBSCAN epsilon (mniejsza = ostrzejsze grupy)
    min_samples: int = 2           # minimalna liczba elementów w klastrze
    max_summary_paragraphs: int = 10
    cache_embeddings: bool = True
    log_level: str = "INFO"

config = Config()

# --------------------
# LOGGING (cross-platform)
# --------------------
def _logs_base() -> Path:
    if platform.system() in ("Windows", "Darwin"):
        return Path.home() / "Documents" / "DocumentAnalysis" / "logs"
    return Path.home() / "DocumentAnalysis" / "logs"

def setup_logging():
    log_dir = _logs_base()
    log_dir.mkdir(parents=True, exist_ok=True)
    logging.basicConfig(
        level=getattr(logging, config.log_level),
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_dir / "analysis.log", encoding='utf-8'),
            logging.StreamHandler()
        ]
    )

setup_logging()
logger = logging.getLogger(__name__)

# --------------------
# OpenAI client (opcjonalny)
# --------------------
client = None
try:
    if config.api_key:
        client = OpenAI(api_key=config.api_key)
        logger.info("✅ OpenAI client initialized")
    else:
        logger.warning("⚠️ No OpenAI API key found. Running in offline mode - summaries unavailable.")
except Exception as e:
    logger.warning(f"⚠️ OpenAI client initialization failed: {e}")

# --------------------
# Embedding model
# --------------------
embedder = None
try:
    logger.info(f"Use pytorch device_name: cpu")
    logger.info(f"Load pretrained SentenceTransformer: {config.embedding_model}")
    embedder = SentenceTransformer(config.embedding_model)
    logger.info(f"✅ Loaded embedding model: {config.embedding_model}")
except Exception as e:
    logger.error(f"❌ Failed to load embedding model: {e}")
    sys.exit(1)

# --------------------
# FILE LOADERS
# --------------------
def load_paragraphs_from_docx(file_path: str) -> List[str]:
    """Load paragraphs from a Word document"""
    try:
        file_path = Path(file_path)
        doc = docx.Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        logger.info(f"📄 Loaded {len(paragraphs)} paragraphs from DOCX: {file_path.name}")
        return paragraphs
    except Exception as e:
        logger.error(f"Error loading DOCX {file_path}: {e}")
        raise

def load_paragraphs_from_pdf(file_path: str) -> List[str]:
    """Load paragraphs from a PDF document (fixed: do not use doc after close)"""
    file_path = Path(file_path)
    try:
        doc = fitz.open(file_path)
        try:
            paragraphs: List[str] = []
            page_count = len(doc)  # liczba stron PRZED zamknięciem
            for page in doc:
                text = page.get_text("text")
                current_para: List[str] = []
                for line in text.split("\n"):
                    line = line.strip()
                    if line:
                        current_para.append(line)
                    elif current_para:
                        paragraphs.append(" ".join(current_para))
                        current_para = []
                if current_para:
                    paragraphs.append(" ".join(current_para))
            logger.info(f"📄 Loaded {len(paragraphs)} paragraphs from PDF: {file_path.name} ({page_count} pages)")
            # Odfiltruj bardzo krótkie ścinki
            return [p for p in paragraphs if len(p) > 20]
        finally:
            doc.close()
    except Exception as e:
        logger.error(f"Error loading PDF {file_path}: {e}")
        raise

def load_paragraphs_from_txt(file_path: str) -> List[str]:
    """Load paragraphs from a text file"""
    file_path = Path(file_path)
    try:
        encodings = ('utf-8', 'utf-8-sig', 'latin1', 'cp1252')
        content = None
        for enc in encodings:
            try:
                content = file_path.read_text(encoding=enc)
                logger.info(f"📄 Successfully read TXT with {enc} encoding")
                break
            except UnicodeDecodeError:
                continue
        if content is None:
            raise ValueError(f"Could not decode TXT: {file_path}")
        # Akapity = bloki rozdzielone pustą linią
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
        logger.info(f"📄 Loaded {len(paragraphs)} paragraphs from TXT: {file_path.name}")
        return paragraphs
    except Exception as e:
        logger.error(f"Error loading TXT {file_path}: {e}")
        raise

def load_paragraphs(file_path: str) -> List[str]:
    """Dispatch loader by extension"""
    p = Path(file_path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {p}")
    if not p.is_file():
        raise ValueError(f"Path is not a file: {p}")
    ext = p.suffix.lower()
    if ext == ".pdf":
        return load_paragraphs_from_pdf(str(p))
    if ext == ".docx":
        return load_paragraphs_from_docx(str(p))
    if ext == ".txt":
        return load_paragraphs_from_txt(str(p))
    raise ValueError(f"Unsupported file format: {ext}. Supported: .pdf, .docx, .txt")

# --------------------
# EMBEDDINGS & CACHING (stabilny hash)
# --------------------
def _stable_hash(texts: List[str], model_id: str) -> str:
    """Stable SHA-256 over model id + paragraphs (order-sensitive)"""
    h = hashlib.sha256()
    h.update(model_id.encode("utf-8"))
    for t in texts:
        h.update(b"\x00")
        h.update(t.encode("utf-8", "ignore"))
    return h.hexdigest()

class EmbeddingCache:
    """Cache embeddings to avoid recomputation (cross-platform path)"""
    def __init__(self, cache_dir: Optional[str] = None):
        if cache_dir:
            self.cache_dir = Path(cache_dir)
        else:
            if platform.system() in ("Windows", "Darwin"):
                self.cache_dir = Path.home() / "Documents" / "DocumentAnalysis" / "embedding_cache"
            else:
                self.cache_dir = Path.home() / "DocumentAnalysis" / "embedding_cache"
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"📦 Cache directory: {self.cache_dir}")

    def get_cache_path(self, key: str) -> Path:
        return self.cache_dir / f"{key}.npy"

    def get(self, texts: List[str], model_id: str) -> Optional[np.ndarray]:
        key = _stable_hash(texts, model_id)
        path = self.get_cache_path(key)
        try:
            if path.exists():
                logger.info("📦 Loading embeddings from cache")
                return np.load(path)
        except Exception as e:
            logger.warning(f"⚠️ Failed to load cache: {e}")
        return None

    def save(self, texts: List[str], embeddings: np.ndarray, model_id: str):
        key = _stable_hash(texts, model_id)
        path = self.get_cache_path(key)
        try:
            np.save(path, embeddings)
            logger.info("💾 Saved embeddings to cache")
        except Exception as e:
            logger.warning(f"⚠️ Failed to save cache: {e}")

cache = EmbeddingCache() if config.cache_embeddings else None

def get_embeddings(paragraphs: List[str]) -> np.ndarray:
    """Generate or retrieve cached embeddings for paragraphs"""
    if not paragraphs:
        return np.array([])
    if not embedder:
        raise RuntimeError("Embedding model not available")

    if cache:
        cached = cache.get(paragraphs, config.embedding_model)
        if cached is not None:
            return cached

    logger.info(f"🔄 Generating embeddings for {len(paragraphs)} paragraphs...")
    try:
        embeddings = embedder.encode(paragraphs, convert_to_tensor=False, show_progress_bar=True)
        if cache:
            cache.save(paragraphs, embeddings, config.embedding_model)
        return embeddings
    except Exception as e:
        logger.error(f"❌ Failed to generate embeddings: {e}")
        raise

# --------------------
# CLUSTERING
# --------------------
def cluster_paragraphs(embeddings: np.ndarray, eps: float = None, min_samples: int = None) -> np.ndarray:
    """Cluster paragraphs using cosine distances + DBSCAN"""
    if embeddings is None or len(embeddings) == 0:
        return np.array([])
    eps = eps if eps is not None else config.eps
    min_samples = min_samples if min_samples is not None else config.min_samples

    logger.info(f"🔍 Clustering with eps={eps}, min_samples={min_samples}")
    dist_matrix = cosine_distances(embeddings)  # NxN – dla bardzo dużych dokumentów może być pamięciożerne
    clustering = DBSCAN(eps=eps, min_samples=min_samples, metric="precomputed")
    labels = clustering.fit_predict(dist_matrix)

    unique_labels = set(labels)
    n_clusters = len(unique_labels) - (1 if -1 in unique_labels else 0)
    n_noise = list(labels).count(-1)
    logger.info(f"📊 Found {n_clusters} clusters, {n_noise} unique paragraphs")
    return labels

def annotate_paragraphs(paragraphs: List[str], labels: np.ndarray) -> List[str]:
    """Prefix paragraphs with cluster label or leave as UNIQUE"""
    annotated: List[str] = []
    for text, label in zip(paragraphs, labels):
        if label == -1:
            annotated.append(text)
        else:
            annotated.append(f"[SIMILAR-{label:02d}] {text}")
    return annotated

# --------------------
# SUMMARIZATION (optional, uses OpenAI if key provided)
# --------------------
def detect_language(text: str) -> str:
    """Very simple PL/EN heuristic"""
    polish_chars = set("ąćęłńóśźżĄĆĘŁŃÓŚŹŻ")
    polish_count = sum(1 for ch in text if ch in polish_chars)
    return "pl" if polish_count > len(text) * 0.01 else "en"

def summarize_group(paragraphs: List[str], label: int, model_name: str = "gpt-4o") -> str:
    """Generate AI summary for a group of similar paragraphs"""
    if not client:
        sample = paragraphs[:3] if len(paragraphs) > 3 else paragraphs
        # pokaż kilka próbek w trybie offline
        return "⚠️ Offline mode - no summary available. Sample paragraphs:\n" + "\n".join(p[:200] for p in sample)

    sample_paragraphs = paragraphs[:config.max_summary_paragraphs]
    combined_text = " ".join(sample_paragraphs)
    language = detect_language(combined_text)
    lang_instruction = "Respond in Polish" if language == "pl" else "Respond in English"

    prompt = f"""
You are analyzing a document for duplicate content. These paragraphs were identified as similar.
{lang_instruction}.

Paragraphs from cluster SIMILAR-{label:02d} ({len(paragraphs)} total):

{chr(10).join(sample_paragraphs)}

Provide a 2-4 sentence summary explaining:
1. The main topic or theme
2. Why these paragraphs are similar
3. The type of information repeated

Summary:
"""
    try:
        response = client.chat.completions.create(
            model=model_name,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=300
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        logger.error(f"❌ Summarization failed: {e}")
        return f"Error generating summary: {str(e)}"

# --------------------
# OUTPUTS
# --------------------
def save_to_docx(annotated_paragraphs: List[str], summaries: Dict[str, str], output_path: str):
    """Save results to a Word document"""
    try:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        d = docx.Document()
        d.add_heading("Document Similarity Analysis", 0)

        # Stats
        d.add_heading("Summary Statistics", 1)
        d.add_paragraph(f"Total paragraphs: {len(annotated_paragraphs)}")
        d.add_paragraph(f"Similar clusters found: {len(summaries)}")
        unique_count = sum(1 for p in annotated_paragraphs if not p.startswith("[SIMILAR-"))
        d.add_paragraph(f"Unique paragraphs: {unique_count}")

        # Annotated doc
        d.add_page_break()
        d.add_heading("Annotated Document", 1)
        for para in annotated_paragraphs:
            p = d.add_paragraph(para)
            if para.startswith("[SIMILAR-") and p.runs:
                # Guard: jeśli nie ma runs, nie zmieniaj koloru
                p.runs[0].font.color.rgb = docx.shared.RGBColor(255, 0, 0)

        # Summaries
        d.add_page_break()
        d.add_heading("Cluster Summaries", 1)
        for group, summary in sorted(summaries.items()):
            d.add_heading(group, level=2)
            d.add_paragraph(summary)

        d.save(output_path)
        logger.info(f"📄 Saved results to {output_path}")
    except Exception as e:
        logger.error(f"❌ Failed to save DOCX: {e}")
        raise

def save_to_csv(annotated_paragraphs: List[str], output_path: str):
    """Save paragraph/group mapping to CSV"""
    try:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, mode="w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow(["Paragraph", "Group", "Length"])
            for para in annotated_paragraphs:
                if para.startswith("[SIMILAR-"):
                    parts = para.split("] ", 1)
                    group = parts[0].replace("[", "")
                    text = parts[1] if len(parts) > 1 else ""
                else:
                    group = "UNIQUE"
                    text = para
                w.writerow([text, group, len(text)])
        logger.info(f"📊 Saved CSV to {output_path}")
    except Exception as e:
        logger.error(f"❌ Failed to save CSV: {e}")
        raise

def save_to_json(annotated_paragraphs: List[str], summaries: Dict[str, str], output_path: str):
    """Save results to JSON (stats + paragraphs + summaries)"""
    try:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        data = {
            "statistics": {
                "total_paragraphs": len(annotated_paragraphs),
                "clusters": len(summaries),
                "unique": sum(1 for p in annotated_paragraphs if not p.startswith("[SIMILAR-"))
            },
            "paragraphs": [],
            "summaries": summaries
        }
        for para in annotated_paragraphs:
            if para.startswith("[SIMILAR-"):
                parts = para.split("] ", 1)
                group = parts[0].replace("[", "")
                text = parts[1] if len(parts) > 1 else ""
            else:
                group = "UNIQUE"
                text = para
            data["paragraphs"].append({
                "text": text,
                "group": group,
                "length": len(text)
            })

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        logger.info(f"📋 Saved JSON to {output_path}")
    except Exception as e:
        logger.error(f"❌ Failed to save JSON: {e}")
        raise

# --------------------
# MAIN PROCESSING
# --------------------
def process_document(
    file_path: str,
    model_name: str = "gpt-4o",
    eps: float = None,
    min_samples: int = None,
    chunk_size: int = None
) -> Tuple[List[str], Dict[str, str]]:
    """
    Process a document to find and summarize similar paragraphs
    """
    logger.info(f"📚 Processing document: {Path(file_path).name}")
    try:
        paragraphs = load_paragraphs(file_path)
        logger.info(f"📄 Total paragraphs: {len(paragraphs)}")
        if len(paragraphs) == 0:
            logger.warning("No paragraphs found in document")
            return [], {}

        embeddings = get_embeddings(paragraphs)
        labels = cluster_paragraphs(embeddings, eps, min_samples)
        annotated = annotate_paragraphs(paragraphs, labels)

        # Zbierz do streszczeń tylko klastry (bez UNIQUE)
        grouped: Dict[int, List[str]] = {}
        for para, label in zip(paragraphs, labels):
            if label != -1:
                grouped.setdefault(label, []).append(para)

        summaries: Dict[str, str] = {}
        for label, group_paragraphs in grouped.items():
            try:
                summary = summarize_group(group_paragraphs, label, model_name)
                summaries[f"SIMILAR-{label:02d}"] = summary
            except Exception as e:
                logger.warning(f"⚠️ Failed to summarize cluster {label}: {e}")
                summaries[f"SIMILAR-{label:02d}"] = f"Summary unavailable: {str(e)}"

        logger.info(f"✅ Processing complete: {len(summaries)} clusters found")
        return annotated, summaries
    except Exception as e:
        logger.error(f"❌ Document processing failed: {e}")
        raise

# --------------------
# PROJECT DIRECTORY HELPER
# --------------------
def get_project_directory(project_name: str) -> Path:
    """Get platform-appropriate project directory"""
    if platform.system() in ("Windows", "Darwin"):
        base_dir = Path.home() / "Documents" / "DocumentAnalysis" / "Projects"
    else:
        base_dir = Path.home() / "DocumentAnalysis" / "Projects"
    project_dir = base_dir / project_name
    project_dir.mkdir(parents=True, exist_ok=True)
    return project_dir
