import os
import re
import random
import pandas as pd
import argparse
import platform
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import List, Dict, Tuple, Optional
import logging
from pathlib import Path

# ==============================
# CONFIGURATION AND SETUP - FIXED for cross-platform and robustness
# ==============================

def setup_logging(project_dir: str) -> None:
    """Setup logging configuration"""
    log_file = os.path.join(project_dir, "extract_examples.log")
    os.makedirs(project_dir, exist_ok=True)
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_file, encoding='utf-8'),
            logging.StreamHandler()
        ]
    )

def parse_arguments() -> argparse.Namespace:
    """Parse command line arguments"""
    parser = argparse.ArgumentParser(description="Extract representative examples from document clusters")
    parser.add_argument("--project", "-p", required=True, help="Project name")
    parser.add_argument("--examples", "-n", type=int, default=10, help="Number of examples per cluster (default: 10)")
    parser.add_argument("--strategy", "-s", choices=['balanced', 'diverse', 'representative'], 
                       default='balanced', help="Sampling strategy (default: balanced)")
    parser.add_argument("--exclude-types", nargs='+', default=['TOC', 'Tabela'], 
                       help="Cluster types to exclude (default: TOC Tabela)")
    parser.add_argument("--min-cluster-size", type=int, default=2, 
                       help="Minimum cluster size to include (default: 2)")
    parser.add_argument("--output-format", choices=['csv', 'docx', 'both'], default='both',
                       help="Output format (default: both)")
    return parser.parse_args()

class ClusterExampleExtractor:
    """Main class for extracting representative examples from clusters"""
    
    def __init__(self, project_name: str, base_dir: str = None):
        self.project_name = project_name
        
        # FIXED: Cross-platform base directory handling
        if base_dir is None:
            if platform.system() == "Windows":
                base_dir = str(Path.home() / "Documents" / "DocumentAnalysis" / "Projects")
            elif platform.system() == "Darwin":  # macOS
                base_dir = str(Path.home() / "Documents" / "DocumentAnalysis" / "Projects")
            else:  # Fallback for other systems
                base_dir = str(Path.home() / "DocumentAnalysis" / "Projects")
        
        self.project_dir = os.path.join(base_dir, project_name)

        # init logging FIRST (FIX)
        setup_logging(self.project_dir)
        self.logger = logging.getLogger(__name__)

        # Prefer clusters_v3 if exists, otherwise fall back to clusters (FIX)
        clusters_v3 = os.path.join(self.project_dir, "clusters_v3")
        clusters_v1 = os.path.join(self.project_dir, "clusters")
        if os.path.exists(clusters_v3):
            self.clusters_folder = clusters_v3
        else:
            self.clusters_folder = clusters_v1

        self.results_file = os.path.join(self.project_dir, "results_doc.csv")
        self.summary_v3_file = os.path.join(self.project_dir, "cluster_summary_v3.csv")
        
        # Output files
        self.output_csv = os.path.join(self.project_dir, "cluster_examples.csv")
        self.output_docx = os.path.join(self.project_dir, "cluster_examples.docx")
        self.output_stats = os.path.join(self.project_dir, "cluster_statistics.csv")
        
        # Load existing data (after logger exists)
        self.cluster_metadata = self._load_cluster_metadata()

    def _load_cluster_metadata(self) -> Dict:
        """Load cluster metadata from existing analysis files (robust to missing cols)"""
        metadata: Dict[str, Dict] = {}
        
        # Load cluster summaries if available (v3)
        if os.path.exists(self.summary_v3_file):
            try:
                df = pd.read_csv(self.summary_v3_file, encoding='utf-8')
                # Expecting columns: Group, Summary
                for _, row in df.iterrows():
                    cluster_id = str(row.get('Group', '')).strip()
                    if cluster_id:
                        md = metadata.setdefault(cluster_id, {})
                        md['summary'] = row.get('Summary', '')
            except Exception as e:
                self.logger.warning(f"Could not load cluster summaries (v3): {e}")
        
        # Load detailed results from results_doc.csv
        if os.path.exists(self.results_file):
            try:
                df = pd.read_csv(self.results_file, encoding='utf-8')
                # Our CSV has columns: Paragraph, Group, Length
                if 'Group' in df.columns:
                    counts = df['Group'].value_counts()
                    for cluster_id, cnt in counts.items():
                        cid = str(cluster_id)
                        md = metadata.setdefault(cid, {})
                        md['paragraph_count'] = int(cnt)
                elif 'Cluster_Label' in df.columns:  # legacy fallback
                    counts = df['Cluster_Label'].value_counts()
                    for cluster_id, cnt in counts.items():
                        cid = str(cluster_id)
                        md = metadata.setdefault(cid, {})
                        md['paragraph_count'] = int(cnt)
                else:
                    self.logger.warning("results_doc.csv has no 'Group' column; metadata will be sparse.")
            except Exception as e:
                self.logger.warning(f"Could not load detailed results: {e}")
                
        return metadata

    def _get_sampling_strategy(self, paragraphs: List[str], n_examples: int, strategy: str) -> List[str]:
        """Apply different sampling strategies to select representative paragraphs"""
        if len(paragraphs) <= n_examples:
            return paragraphs[:]
            
        if strategy == 'balanced':
            # Take first, last, and evenly distributed middle paragraphs
            examples = []
            if n_examples >= 1:
                examples.append(paragraphs[0])
            if n_examples >= 2:
                examples.append(paragraphs[-1])
            remaining = max(0, n_examples - len(examples))
            if remaining > 0 and len(paragraphs) > 2:
                middle = paragraphs[1:-1]
                step = max(1, len(middle) // remaining)
                examples.extend(middle[::step][:remaining])
                
        elif strategy == 'diverse':
            indices = []
            step = max(1, int(len(paragraphs) / n_examples))
            for i in range(n_examples):
                idx = min(i * step, len(paragraphs) - 1)
                indices.append(idx)
            examples = [paragraphs[i] for i in sorted(set(indices))]
            
        elif strategy == 'representative':
            examples = [paragraphs[0]]
            if n_examples > 1:
                examples.append(paragraphs[-1])
            if n_examples > 2 and len(paragraphs) > 2:
                middle = paragraphs[1:-1]
                remaining = min(n_examples - 2, len(middle))
                if remaining > 0:
                    examples.extend(random.sample(middle, remaining))
        else:
            examples = random.sample(paragraphs, n_examples)
            
        return examples[:n_examples]

    def _normalize_txt_lines(self, lines: List[str]) -> List[str]:
        """
        Convert lines from cluster txt to clean paragraphs.
        - skip headers starting with '#'
        - strip numbering like '1. ' at line start
        - drop empty lines
        """
        out = []
        for line in lines:
            s = line.strip()
            if not s:
                continue
            if s.startswith("#"):
                continue
            # strip list numbering "1. text"
            s = re.sub(r'^\s*\d+\.\s+', '', s)
            out.append(s)
        return out

    def extract_examples_from_file(self, filepath: str, n_examples: int, strategy: str = 'balanced') -> Tuple[List[str], Dict]:
        """Extract representative paragraphs from a cluster file with statistics"""
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                raw_lines = f.readlines()
        except UnicodeDecodeError:
            with open(filepath, "r", encoding="cp1252") as f:
                raw_lines = f.readlines()
        except Exception as e:
            self.logger.error(f"Error reading file {filepath}: {e}")
            return [], {}

        paragraphs = self._normalize_txt_lines(raw_lines)

        examples = self._get_sampling_strategy(paragraphs, n_examples, strategy)
        
        stats = {
            'total_paragraphs': len(paragraphs),
            'extracted_examples': len(examples),
            'avg_paragraph_length': (sum(len(p) for p in paragraphs) / len(paragraphs)) if paragraphs else 0,
            'coverage_ratio': (len(examples) / len(paragraphs)) if paragraphs else 0
        }
        
        return examples, stats

    def should_include_cluster(self, cluster_name: str, exclude_types: List[str], min_size: int) -> bool:
        """Determine if cluster should be included based on filters"""
        info = self.cluster_metadata.get(cluster_name, {})
        # type info may be absent; treat as Content
        cluster_type = info.get('type', 'Content')
        if cluster_type in exclude_types:
            return False

        cluster_size = info.get('paragraph_count', None)
        if cluster_size is not None and cluster_size < min_size:
            return False
                
        return True

    def create_detailed_docx(self, data: List[Dict], statistics: Dict) -> None:
        """Create a comprehensive Word document with examples and metadata"""
        doc = Document()
        
        title = doc.add_heading('Representative Examples from Clusters', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Project info
        doc.add_heading('Project Information', level=1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        info_data = [
            ['Project', self.project_name],
            ['Number of Clusters', str(len(statistics))],
            ['Total Examples', str(len(data))],
            ['Generated', pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')]
        ]
        for i, (key, value) in enumerate(info_data):
            info_table.cell(i, 0).text = key
            info_table.cell(i, 1).text = value

        # Cluster statistics
        doc.add_heading('Cluster Statistics', level=1)
        if statistics:
            stats_df = pd.DataFrame.from_dict(statistics, orient='index')
            stats_table = doc.add_table(rows=len(stats_df) + 1, cols=len(stats_df.columns) + 1)
            stats_table.style = 'Table Grid'
            stats_table.cell(0, 0).text = 'Cluster'
            for j, col in enumerate(stats_df.columns):
                stats_table.cell(0, j + 1).text = col
            for i, (cluster, row) in enumerate(stats_df.iterrows()):
                stats_table.cell(i + 1, 0).text = str(cluster)
                for j, value in enumerate(row):
                    stats_table.cell(i + 1, j + 1).text = str(value)

        # Examples by cluster
        doc.add_heading('Examples by Cluster', level=1)
        current_cluster = None
        for item in data:
            cluster = item['Cluster']
            if cluster != current_cluster:
                current_cluster = cluster
                doc.add_heading(f"Cluster: {cluster}", level=2)
                # Optional metadata
                if cluster in self.cluster_metadata:
                    metadata = self.cluster_metadata[cluster]
                    meta_text = []
                    meta_text.append(f"Paragraph Count: {metadata.get('paragraph_count', 'N/A')}")
                    if 'summary' in metadata and metadata['summary']:
                        meta_text.append(f"Summary: {metadata['summary']}")
                    doc.add_paragraph(" | ".join(meta_text))

            para = doc.add_paragraph(item['Paragraph'])
            para.style = 'Quote'

        doc.save(self.output_docx)

    def run(self, n_examples: int = 10, strategy: str = 'balanced', 
            exclude_types: List[str] = None, min_cluster_size: int = 2,
            output_format: str = 'both') -> None:
        """Main execution method"""
        
        if exclude_types is None:
            exclude_types = ['TOC', 'Tabela']
            
        self.logger.info(f"Starting example extraction for project: {self.project_name}")
        self.logger.info(f"Clusters folder: {self.clusters_folder}")
        self.logger.info(f"Strategy: {strategy}, Examples per cluster: {n_examples}")
        self.logger.info(f"Excluding types: {exclude_types}, Min cluster size: {min_cluster_size}")
        
        if not os.path.exists(self.clusters_folder):
            self.logger.error(f"Clusters folder not found: {self.clusters_folder}")
            return

        data = []
        statistics = {}
        processed_clusters = 0
        excluded_clusters = 0

        # Process each cluster file
        cluster_files = [f for f in os.listdir(self.clusters_folder) 
                         if f.endswith(".txt") and f.startswith("SIMILAR-")]
        
        self.logger.info(f"Found {len(cluster_files)} cluster files")

        for filename in sorted(cluster_files):
            cluster_name = filename.replace(".txt", "")
            
            if not self.should_include_cluster(cluster_name, exclude_types, min_cluster_size):
                excluded_clusters += 1
                self.logger.debug(f"Excluding cluster {cluster_name} (filtered out)")
                continue
                
            filepath = os.path.join(self.clusters_folder, filename)
            examples, stats = self.extract_examples_from_file(filepath, n_examples, strategy)
            
            if not examples:
                self.logger.warning(f"No examples extracted from {cluster_name}")
                continue

            for ex in examples:
                data.append({
                    "Cluster": cluster_name,
                    "Paragraph": ex,
                    "Character_Count": len(ex),
                    "Word_Count": len(ex.split())
                })
            
            statistics[cluster_name] = stats
            processed_clusters += 1
            
            self.logger.info(f"Extracted {len(examples)} examples from {cluster_name} "
                             f"({stats['total_paragraphs']} total paragraphs)")

        if not data:
            self.logger.error("No data extracted. Check your filters and cluster files.")
            return

        if output_format in ['csv', 'both']:
            df = pd.DataFrame(data)
            df.to_csv(self.output_csv, index=False, encoding="utf-8")
            self.logger.info(f"CSV results saved: {self.output_csv}")

        if output_format in ['docx', 'both']:
            self.create_detailed_docx(data, statistics)
            self.logger.info(f"DOCX document saved: {self.output_docx}")

        if statistics:
            stats_df = pd.DataFrame.from_dict(statistics, orient='index')
            stats_df.to_csv(self.output_stats, encoding="utf-8")
            self.logger.info(f"Statistics saved: {self.output_stats}")

        self.logger.info("\n" + "="*50)
        self.logger.info("EXTRACTION COMPLETED!")
        self.logger.info(f"Project: {self.project_name}")
        self.logger.info(f"Processed clusters: {processed_clusters}")
        self.logger.info(f"Excluded clusters: {excluded_clusters}")
        self.logger.info(f"Total examples: {len(data)}")
        self.logger.info(f"Average examples per cluster: {len(data)/processed_clusters:.1f}")
        self.logger.info("="*50)


# ==============================
# MAIN PROGRAM
# ==============================
def main():
    """Main execution function"""
    args = parse_arguments()
    try:
        extractor = ClusterExampleExtractor(args.project)
        extractor.run(
            n_examples=args.examples,
            strategy=args.strategy,
            exclude_types=args.exclude_types,
            min_cluster_size=args.min_cluster_size,
            output_format=args.output_format
        )
    except Exception as e:
        logging.error(f"Fatal error during execution: {e}")
        raise


if __name__ == "__main__":
    random.seed(42)
    main()
